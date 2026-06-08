"""Gerador de DOCX de fichas de critério via python-docx.

A intenção: o servidor que abre o DOCX em Word pode copiar trechos
DIRETAMENTE para o seu TR/ETP sem retrabalho de formatação. Por isso
usamos Verdana 10pt como base (fonte de sistema definida no Manual
GESP v1.6 para uso documental), e títulos em estilo Heading 1/2.
"""

from __future__ import annotations

import re
from io import BytesIO

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

PRIMARY_RED = RGBColor(0xED, 0x1C, 0x24)


def _set_default_font(doc: Document, font_name: str = "Verdana", size_pt: int = 10) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)


def _strip_html(html: str) -> str:
    """Remoção mínima de tags HTML.

    A maioria dos campos chega como rich-text leve (parágrafos, listas,
    negritos). Para v1.0 fazemos uma limpeza simples — em uma iteração
    posterior podemos integrar bleach + python-docx mais granular.
    """
    if not html:
        return ""
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"</p>\s*<p[^>]*>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def _adicionar_paragrafos(doc: Document, conteudo_html: str) -> None:
    texto = _strip_html(conteudo_html)
    for paragrafo in texto.split("\n\n"):
        if not paragrafo.strip():
            continue
        p = doc.add_paragraph(paragrafo.strip())
        p.paragraph_format.space_after = Pt(6)


def gerar_docx_criterio(criterio) -> bytes:
    """Renderiza a ficha do critério em DOCX e devolve bytes."""
    doc = Document()
    _set_default_font(doc, "Verdana", 10)

    # Cabeçalho institucional
    header = doc.sections[0].header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run(
        "Governo do Estado de São Paulo — Plataforma Estadual de "
        "Sustentabilidade em Contratações Públicas"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = PRIMARY_RED

    # Rodapé com URL e versão
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cfg = settings.PESCP
    url_publica = cfg.get("PUBLIC_URL", "") + criterio.get_absolute_url()
    footer_run = footer_p.add_run(f"PESCP v{cfg.get('VERSAO_ROTULO', '1.0')} — {url_publica}")
    footer_run.font.size = Pt(8)

    # ---------- Título ----------
    titulo = doc.add_heading(criterio.titulo, level=1)
    titulo.runs[0].font.size = Pt(16)
    titulo.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    meta = doc.add_paragraph()
    meta.add_run(f"Código: {criterio.codigo}").bold = True
    meta.add_run(
        f"   |   Objeto: {criterio.objeto_contratual}"
        f"   |   Fase: {criterio.get_fase_processo_display()}"
        f"   |   Nível: {criterio.get_nivel_ambicao_display()}"
    )

    # ---------- Texto da Providência (destaque) ----------
    doc.add_heading("Texto pronto para inserção", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(_strip_html(criterio.texto_providencia))
    run.font.size = Pt(11)

    # ---------- Demais seções ----------
    secoes = [
        ("Determinações Principais", criterio.determinacoes_principais),
        ("Justificativa Técnica", criterio.justificativa_tecnica),
        ("Método de Verificação", criterio.metodo_verificacao),
    ]
    if criterio.precaucoes:
        secoes.append(("Precauções e Ressalvas", criterio.precaucoes))

    for titulo_secao, conteudo in secoes:
        doc.add_heading(titulo_secao, level=2)
        _adicionar_paragrafos(doc, conteudo)

    # ---------- Referências ----------
    doc.add_heading("Referências", level=2)

    if criterio.ods.exists():
        p = doc.add_paragraph()
        p.add_run("ODS aplicáveis: ").bold = True
        p.add_run(", ".join(str(o) for o in criterio.ods.all()))

    if criterio.normas.exists():
        p = doc.add_paragraph()
        p.add_run("Normas legais: ").bold = True
        p.add_run("; ".join(n.titulo_curto for n in criterio.normas.all()))

    if criterio.selos.exists():
        p = doc.add_paragraph()
        p.add_run("Selos / certificações: ").bold = True
        p.add_run("; ".join(s.nome for s in criterio.selos.all()))

    if criterio.fonte:
        p = doc.add_paragraph()
        p.add_run("Fonte: ").bold = True
        p.add_run(criterio.fonte)

    p = doc.add_paragraph()
    p.add_run(
        "\nEste documento foi gerado pela Plataforma Estadual de Sustentabilidade "
        "em Contratações Públicas (PESCP). Conteúdo em validação contínua."
    ).italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
